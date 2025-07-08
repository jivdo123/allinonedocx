import os
import logging
import re
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from docx import Document
from docx.shared import Pt

# --- Configuration ---
# IMPORTANT: Replace "YOUR_TELEGRAM_BOT_TOKEN" with your actual bot token.
# It is highly recommended to use environment variables for security.
TELEGRAM_BOT_TOKEN = "8127720127:AAFeFVi4a2ZXmY-osUz9HjreJT4ZCfe4mtc"
TEMP_DIR = "temp_docx_files"
MAX_TABLES_PER_FILE = 30

# --- Logging Setup ---
# This helps in debugging by printing informative messages to your console.
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# --- In-memory storage for files waiting for the /convert command ---
# Structure: {chat_id: [file_path1, file_path2, ...]}
user_files_queue = {}


# --- Core Logic: Parsing and Generation Functions ---

def parse_plain_text_file(file_path):
    """
    Parses a plain text .docx file and extracts question data into a structured list.

    Args:
        file_path (str): The path to the .docx file.

    Returns:
        tuple: A list of question data dictionaries and a count of skipped blocks.
    """
    try:
        doc = Document(file_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        logger.error(f"Error reading docx file {file_path}: {e}")
        return [], 0

    # This regex splits the text into blocks, keeping the delimiter (e.g., "Q1.", "20.")
    question_blocks = re.split(r'\n(?=Q\d*\.\s|\d+\.\s)', full_text.strip())
    
    extracted_data = []
    skipped_blocks = 0

    for block in question_blocks:
        if not block.strip():
            continue

        lines = [line.strip() for line in block.strip().split('\n') if line.strip()]
        
        # Validate that the block has the minimum required lines
        if len(lines) < 6:
            skipped_blocks += 1
            logger.warning(f"Skipping block with less than 6 lines: {block}")
            continue

        try:
            # Extract Question (removing prefix like "Q1." or "1.")
            question_text = re.sub(r'^(Q\d*\.|\d+\.)\s*', '', lines[0])

            # Extract Options (removing prefixes like "a.", "b.")
            options = [re.sub(r'^[a-d]\.\s*', '', line) for line in lines[1:5]]

            # Extract Correct Answer ID (e.g., 'b' from "Correct Option: b")
            correct_option_line = lines[5]
            match = re.search(r':\s*([a-d])', correct_option_line, re.IGNORECASE)
            if not match:
                skipped_blocks += 1
                logger.warning(f"Could not find correct option in line: {correct_option_line}")
                continue
            correct_id = match.group(1).lower()

            # --- MODIFIED LOGIC FOR MULTI-LINE EXPLANATION ---
            explanation = ""
            # Check if there are any lines after the "Correct Option" line.
            if len(lines) > 6:
                # The rest of the lines belong to the explanation.
                explanation_lines = lines[6:]
                
                # The first line of the explanation might have the "Explanation:" prefix. Remove it.
                if explanation_lines and explanation_lines[0].lower().startswith('explanation:'):
                    explanation_lines[0] = re.sub(r'^explanation:\s*', '', explanation_lines[0], flags=re.IGNORECASE).strip()
                
                # Join the lines (including the modified first line) to form the full explanation.
                explanation = "\n".join(explanation_lines)
            # --- END OF MODIFIED LOGIC ---

            extracted_data.append({
                "question": question_text,
                "options": options,
                "correct_id": correct_id,
                "explanation": explanation
            })
        except (IndexError, Exception) as e:
            skipped_blocks += 1
            logger.warning(f"Skipping malformed block. Error: {e}. Block: {block}")

    return extracted_data, skipped_blocks

def parse_table_docx_file(file_path):
    """
    Parses a .docx file with existing tables and extracts question data.

    Args:
        file_path (str): The path to the .docx file.

    Returns:
        list: A list of question data dictionaries.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        logger.error(f"Error reading docx file {file_path}: {e}")
        return []

    extracted_data = []
    for table in doc.tables:
        # Validate table structure
        if not (len(table.rows) >= 8 and len(table.columns) >= 3):
            logger.warning("Skipping a table with incorrect dimensions.")
            continue
        
        try:
            question_text = table.cell(0, 1).text.strip()
            options = [table.cell(i, 1).text.strip() for i in range(2, 6)]
            explanation = table.cell(6, 1).text.strip()
            
            correct_id = ''
            option_chars = ['a', 'b', 'c', 'd']
            # Find which option is marked 'correct'
            for i in range(2, 6):
                if table.cell(i, 2).text.strip().lower() == 'correct':
                    correct_id = option_chars[i-2]
                    break
            
            if not correct_id:
                logger.warning(f"Could not find a 'correct' answer for question: {question_text[:50]}...")
                continue

            extracted_data.append({
                "question": question_text,
                "options": options,
                "correct_id": correct_id,
                "explanation": explanation
            })
        except IndexError:
            logger.warning(f"Skipping a malformed table in {file_path}")
            continue
    
    return extracted_data


def create_output_docs(data_list, chat_id):
    """
    Generates paginated .docx files from a list of question data.

    Args:
        data_list (list): A list of question data dictionaries.
        chat_id (int): The user's chat ID for naming files.

    Returns:
        list: A list of file paths for the generated .docx files.
    """
    if not data_list:
        return []

    output_files = []
    table_count = 0
    doc = Document()

    for i, q_data in enumerate(data_list):
        # Paginate: create a new file after 30 tables
        if table_count > 0 and table_count % MAX_TABLES_PER_FILE == 0:
            part = len(output_files) + 1
            output_path = os.path.join(TEMP_DIR, f"{chat_id}_output_part_{part}.docx")
            doc.save(output_path)
            output_files.append(output_path)
            doc = Document() # Start a new document
        
        # Create and style the table
        table = doc.add_table(rows=8, cols=3)
        table.style = 'Table Grid'
        
        # --- Column 1: Static Labels ---
        static_labels = ["Question", "Type", "Option", "Option", "Option", "Option", "Solution", "Marks"]
        for row_idx, label in enumerate(static_labels):
            table.cell(row_idx, 0).text = label

        # --- Row 1: Question ---
        cell_q_text = table.cell(0, 1)
        cell_q_text.merge(table.cell(0, 2))
        run_q = cell_q_text.paragraphs[0].add_run(q_data['question'])
        run_q.font.size = Pt(14)

        # --- Row 2: Type ---
        cell_type_text = table.cell(1, 1)
        cell_type_text.merge(table.cell(1, 2))
        cell_type_text.text = "multiple_choice"

        # --- Rows 3-6: Options ---
        option_chars = ['a', 'b', 'c', 'd']
        for row_idx, option_text in zip(range(2, 6), q_data['options']):
            run_opt = table.cell(row_idx, 1).paragraphs[0].add_run(option_text)
            run_opt.font.size = Pt(12)
            is_correct = (option_chars[row_idx-2] == q_data['correct_id'])
            table.cell(row_idx, 2).text = "correct" if is_correct else "incorrect"

        # --- Row 7: Solution/Explanation ---
        cell_sol_text = table.cell(6, 1)
        cell_sol_text.merge(table.cell(6, 2))
        run_sol = cell_sol_text.paragraphs[0].add_run(q_data['explanation'])
        run_sol.font.size = Pt(12)
        
        # --- Row 8: Marks ---
        table.cell(7, 1).text = "4"
        table.cell(7, 2).text = "1"
        
        table_count += 1
        doc.add_paragraph() # Add spacing between tables

    # Save the final (or only) document
    if table_count > 0:
        part = len(output_files) + 1
        output_path = os.path.join(TEMP_DIR, f"{chat_id}_output_part_{part}.docx")
        doc.save(output_path)
        output_files.append(output_path)

    return output_files

# --- Bot Command and Message Handlers ---

async def start_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Sends a welcome message for the /start command."""
    await update.message.reply_text(
        "Hello! I am the Docx Question Formatter Bot. 🤖\n\n"
        "Send me a `.docx` file, and I'll handle the rest. "
        "Use /help to see detailed instructions."
    )

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Sends a detailed help message."""
    await update.message.reply_text(
        "--- How to Use ---\n\n"
        "🔹 **Workflow 1: Plain Text to Tables**\n"
        "1. Create a `.docx` file.\n"
        "2. Write questions in this format (minimum 6 lines per question):\n"
        "   Q1. What is 2+2?\n"
        "   a. 3\n"
        "   b. 4\n"
        "   c. 5\n"
        "   d. 6\n"
        "   Correct Option: b\n"
        "   Explanation: This is an optional explanation.\n"
        "   It can even span multiple lines.\n"
        "3. Send the file to me. I will process it instantly.\n\n"
        "🔹 **Workflow 2: Merge & Reformat Existing Tables**\n"
        "1. Send me one or more `.docx` files that already contain question tables.\n"
        "2. I will confirm each file receipt.\n"
        "3. When you've sent all files, type the `/convert` command.\n"
        "4. I will merge, reformat, and send back the results.\n\n"
        "Output is paginated into a new file every 30 questions."
    )

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handles all incoming documents."""
    message = update.message
    if not message.document or not message.document.file_name.endswith('.docx'):
        await message.reply_text("Error: Please send a valid `.docx` file. 🚫")
        return

    chat_id = message.chat_id
    
    try:
        file = await context.bot.get_file(message.document.file_id)
        
        # Create a temporary directory for the user if it doesn't exist
        user_temp_dir = os.path.join(TEMP_DIR, str(chat_id))
        os.makedirs(user_temp_dir, exist_ok=True)
        file_path = os.path.join(user_temp_dir, f"{file.file_unique_id}.docx")
        await file.download_to_drive(file_path)
        
        # --- Analyze file to determine workflow ---
        await message.reply_text("Analyzing file... 🧐")
        
        try:
            doc = Document(file_path)
            is_table_doc = bool(doc.tables)
        except Exception:
            await message.reply_text("Error: This `.docx` file appears to be corrupted.")
            os.remove(file_path)
            return

        if is_table_doc:
            # --- Workflow 2: Queue the file ---
            user_files_queue.setdefault(chat_id, []).append(file_path)
            await message.reply_text(
                f"✅ File '{message.document.file_name}' (containing tables) received.\n\n"
                "Send more files or use the /convert command to finish."
            )
        else:
            # --- Workflow 1: Process the plain text file immediately ---
            await message.reply_text("Plain text file detected. Processing now... ⚙️")
            
            data, skipped_count = parse_plain_text_file(file_path)
            os.remove(file_path) # Clean up downloaded file immediately
            
            if not data:
                await message.reply_text("Could not find any valid questions. Please check the file's formatting and use /help for an example.")
                return

            output_files = create_output_docs(data, chat_id)
            
            await message.reply_text(f"Processing complete. Found {len(data)} questions. Sending back the formatted file(s)...")
            
            for doc_path in output_files:
                await context.bot.send_document(chat_id=chat_id, document=open(doc_path, 'rb'))
                os.remove(doc_path) # Clean up generated file after sending

            summary_message = f"Successfully converted {len(data)} questions."
            if skipped_count > 0:
                summary_message += f"\nSkipped {skipped_count} blocks due to formatting errors."
            await message.reply_text(summary_message)

    except Exception as e:
        logger.error(f"Error in handle_document for user {chat_id}: {e}", exc_info=True)
        await message.reply_text("An unexpected error occurred. Please try again. 🙏")


async def convert_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Processes all queued files for a user on /convert command."""
    chat_id = update.message.chat_id

    if chat_id not in user_files_queue or not user_files_queue[chat_id]:
        await update.message.reply_text(
            "There are no files to convert. 🤔 Please upload `.docx` files with tables first."
        )
        return
        
    await update.message.reply_text(
        f"Starting conversion for the {len(user_files_queue[chat_id])} file(s) you sent. This may take a moment... ⏳"
    )

    all_data = []
    queued_files = user_files_queue[chat_id]
    for file_path in queued_files:
        all_data.extend(parse_table_docx_file(file_path))

    if not all_data:
        await update.message.reply_text("Could not extract any valid question tables from the files provided.")
    else:
        output_files = create_output_docs(all_data, chat_id)
        await update.message.reply_text(f"Conversion complete. Found {len(all_data)} total questions. Sending back the formatted file(s)...")
        for doc_path in output_files:
            await context.bot.send_document(chat_id=chat_id, document=open(doc_path, 'rb'))
            os.remove(doc_path)

    # --- Cleanup: remove all processed files for the user ---
    for file_path in queued_files:
        try:
            os.remove(file_path)
        except OSError as e:
            logger.error(f"Error cleaning up queued file {file_path}: {e}")
    del user_files_queue[chat_id]
    await update.message.reply_text("All done! Your queue has been cleared. ✨")


# --- Main Bot Execution ---
def main():
    """Starts the bot and sets up handlers."""
    # Create the temporary directory for file storage if it doesn't exist
    os.makedirs(TEMP_DIR, exist_ok=True)

    if TELEGRAM_BOT_TOKEN == "YOUR_TELEGRAM_BOT_TOKEN":
        print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
        print("!!! ERROR: Please replace 'YOUR_TELEGRAM_BOT_TOKEN' in the   !!!")
        print("!!! script with your bot's actual token from BotFather.     !!!")
        print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
        return

    application = Application.builder().token(TELEGRAM_BOT_TOKEN).build()

    # Add handlers for commands
    application.add_handler(CommandHandler("start", start_command))
    application.add_handler(CommandHandler("help", help_command))
    application.add_handler(CommandHandler("convert", convert_command))

    # Add a handler for receiving documents
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))

    # Start the bot
    print("Bot is running... Press Ctrl-C to stop.")
    application.run_polling()

if __name__ == '__main__':
    main()
            
